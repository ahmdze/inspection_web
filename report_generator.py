import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def set_rtl_and_justify(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = align
    pPr = paragraph._element.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))


def set_font_style(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    for attr, value in [
        ("w:ascii", "Times New Roman"),
        ("w:hAnsi", "Times New Roman"),
        ("w:cs", "Times New Roman"),
    ]:
        rFonts.set(qn(attr), value)
    rPr.append(rFonts)
    rPr.append(OxmlElement("w:rtl"))
    if bold:
        rPr.append(OxmlElement("w:bCs"))


def clean_numeric_value(value):
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return "" if value is None else str(value)


def format_date_only(value):
    if not value:
        return ""
    try:
        if isinstance(value, datetime):
            return value.strftime("%d/%m/%Y")
        return datetime.strptime(str(value).strip(), "%Y-%m-%d").strftime("%d/%m/%Y")
    except Exception:
        return str(value)


def _add_title(doc, text, size, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    set_rtl_and_justify(p, align)
    set_font_style(p.add_run(text), size=size, bold=True)
    return p


def _add_field(doc, label, value):
    p = doc.add_paragraph()
    set_rtl_and_justify(p)
    run_label = p.add_run(f"{label}: ")
    set_font_style(run_label, size=12, bold=True)
    run_value = p.add_run(clean_numeric_value(value))
    set_font_style(run_value, size=12)
    return p


def _add_value_only(doc, value):
    """فقرة تحتوي على القيمة بدون عنوان الحقل"""
    p = doc.add_paragraph()
    set_rtl_and_justify(p)
    run_value = p.add_run(clean_numeric_value(value))
    set_font_style(run_value, size=12)
    return p


# ─── دالة تكرارية لرسم شجرة الأقسام ─────────────────────────────────────────
def _render_section_tree(doc, flat_sections, section_id, level=0):
    """
    تعمل بشكل تكراري على أي عمق من الأقسام.

    flat_sections: قاموس  { section_id(int) : { id, name, order, parent_id, data:[...] } }

    level=0  → قسم رئيسي        (حجم خط 16)
    level=1  → قسم فرعي         (حجم خط 14)
    level=2+ → قسم فرعي فرعي   (حجم خط 13)

    بيانات القسم الرئيسي تُعرض بـ label+value.
    بيانات الأقسام الفرعية تُعرض بـ value فقط.
    """
    section = flat_sections.get(section_id)
    if not section:
        return

    # ─ تحقق من وجود بيانات في هذا الفرع أو أي تابع له ─
    def has_data(sid):
        sec = flat_sections.get(sid)
        if not sec:
            return False
        if sec.get("data"):
            return True
        children_ids = [s["id"] for s in flat_sections.values() if s.get("parent_id") == sid]
        return any(has_data(cid) for cid in children_ids)

    if not has_data(section_id):
        return

    # ─ حجم الخط حسب المستوى ─
    title_size = 16 if level == 0 else (14 if level == 1 else 13)

    # ─ كتابة اسم القسم ─
    _add_title(doc, section.get("name", ""), title_size)

    # ─ كتابة بيانات هذا القسم ─
    section_data = sorted(section.get("data", []), key=lambda x: x.get("order", 0))
    for item in section_data:
        if level == 0:
            # القسم الرئيسي: اسم الحقل + القيمة
            _add_field(doc, item.get("label", ""), item.get("value", ""))
        else:
            # الأقسام الفرعية بأي مستوى: القيمة فقط
            _add_value_only(doc, item.get("value", ""))

    # ─ الأقسام الأبناء بشكل تكراري ─
    children = sorted(
        [s for s in flat_sections.values() if s.get("parent_id") == section_id],
        key=lambda x: x.get("order", 0),
    )
    for child in children:
        _render_section_tree(doc, flat_sections, child["id"], level + 1)


# ─── بناء قاموس مفلطح من بيانات الإجابات ────────────────────────────────────
def build_flat_sections(all_sections_db, all_submissions, field_map):
    """
    تحوّل قائمة الأقسام من قاعدة البيانات + الإجابات إلى قاموس مفلطح.

    all_sections_db : نتيجة db.query(Section).all()
    all_submissions : نتيجة db.query(Submission).filter(...).all()
    field_map       : { field_key: FormField }

    تُرجع: { section_id(int): { id, name, order, parent_id, data:[{label,value,order}] } }
    """
    import json

    flat = {}
    for sec in all_sections_db:
        flat[sec.id] = {
            "id": sec.id,
            "name": sec.name,
            "order": sec.order,
            "parent_id": sec.parent_id,
            "data": [],
        }

    for sub in all_submissions:
        try:
            answers = json.loads(sub.answers_json)
        except Exception:
            continue

        for key, val in answers.items():
            if not val or key.startswith("rec_enable_") or key.startswith("rec_"):
                continue
            field = field_map.get(key)
            if not field or field.section_id is None:
                continue
            if field.section_id in flat:
                flat[field.section_id]["data"].append(
                    {
                        "label": field.label,
                        "value": str(val),
                        "order": field.order,
                    }
                )

    return flat


# ─── الدالة الرئيسية لتوليد التقرير ─────────────────────────────────────────
def build_web_report(data: dict, output_folder: str) -> str:
    """
    data يحتوي على:
      general               : { institution, visit_date, ...حقول عامة }
      flat_sections         : { section_id(int): { id, name, order, parent_id, data:[...] } }
      recommendations       : { rec_a: [...], rec_b: [...], rec_c: [...], rec_d: [...] }
      recommendation_categories : [{ key, label, order }, ...]
    """
    general = data.get("general", {})
    inst = general.get("institution") or "غير محدد"
    date = general.get("visit_date") or "غير محدد"
    safe_name = str(inst).replace("/", "-").replace("\\", "-")
    safe_date = str(date).replace("/", "-").replace("\\", "-")
    path = os.path.join(output_folder, f"تقرير_{safe_name}_{safe_date}.docx")

    doc = Document()

    # ─── عنوان التقرير ───────────────────────────────────────────────
    p_subj = doc.add_paragraph()
    set_rtl_and_justify(p_subj, WD_ALIGN_PARAGRAPH.CENTER)
    set_font_style(p_subj.add_run("م/ زيارة تفتيشية"), size=18, bold=True)

    # ─── المقدمة ─────────────────────────────────────────────────────
    intro_text = (
        f"استناداً إلى الخطة السنوية لشعبة تفتيش المؤسسات الصحية الحكومية، "
        f"أجرى فريق من قسم التفتيش زيارة تفتيشية الى ({inst}) بتاريخ ({format_date_only(date)})"
    )
    p_intro = doc.add_paragraph()
    set_rtl_and_justify(p_intro)
    p_intro.paragraph_format.line_spacing = 1.5
    set_font_style(p_intro.add_run(intro_text), size=12)

    p_note = doc.add_paragraph()
    set_rtl_and_justify(p_note)
    set_font_style(p_note.add_run("وتم ملاحظة الاتي :"), size=12, bold=True)
    doc.add_paragraph(" ")

    # ─── المعلومات العامة ────────────────────────────────────────────
    for label, value in general.items():
        if label in ["institution", "visit_date"]:
            continue
        _add_field(doc, label, value)
    doc.add_paragraph(" ")

    # ─── الأقسام بالهيكل التكراري ───────────────────────────────────
    flat_sections = data.get("flat_sections", {})

    if flat_sections:
        # الأقسام الجذرية (بدون أب)
        root_sections = sorted(
            [s for s in flat_sections.values() if s.get("parent_id") is None],
            key=lambda x: x.get("order", 0),
        )
        for root in root_sections:
            _render_section_tree(doc, flat_sections, root["id"], level=0)
            doc.add_paragraph(" ")

    else:
        # ─ توافق مع الهيكل القديم (مستويان فقط) ─
        sections = data.get("sections", {})
        for _, section in sorted(sections.items(), key=lambda item: item[1].get("order", 0)):
            section_data = sorted(section.get("data", []), key=lambda x: x.get("order", 0))
            subsections = sorted(
                section.get("subsections", {}).items(),
                key=lambda item: item[1].get("order", 0),
            )
            has_any = bool(section_data) or any(s.get("data") for _, s in subsections)
            if not has_any:
                continue
            _add_title(doc, section.get("name", "محور"), 16)
            for item in section_data:
                _add_field(doc, item.get("label", ""), item.get("value", ""))
            for _, sub in subsections:
                sub_data = sorted(sub.get("data", []), key=lambda x: x.get("order", 0))
                if not sub_data:
                    continue
                _add_title(doc, sub.get("name", ""), 14)
                for item in sub_data:
                    _add_value_only(doc, item.get("value", ""))
            doc.add_paragraph(" ")

    # ─── التوصيات ────────────────────────────────────────────────────
    recommendations = data.get("recommendations", {})
    rec_categories = data.get("recommendation_categories") or [
        {"key": "rec_a", "label": "أ/ الإيعاز إلى دائرة صحة بغداد الرصافة/ قسم التخطيط:"},
        {"key": "rec_b", "label": "ب/ الإيعاز إلى شعبة التحقيقات/ قسمنا، بتشكيل لجنة تحقيقية بخصوص:"},
        {"key": "rec_c", "label": "ج/ الإيعاز إلى إدارة المستشفى بخصوص:"},
        {"key": "rec_d", "label": "د/ أخرى:"},
    ]

    has_recommendations = any(
        recommendations.get(cat["key"]) for cat in rec_categories
    )

    if has_recommendations:
        # صفحة جديدة
        doc.add_page_break()
        # عنوان التوصيات
        _add_title(doc, "التوصيات:", 16, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph(" ")

        for cat in rec_categories:
            items = [i for i in recommendations.get(cat["key"], []) if str(i).strip()]
            if not items:
                continue
            # العنوان الفرعي للفئة
            _add_title(doc, cat["label"], 13)
            # الإجابات مرقمة
            for idx, item in enumerate(items, start=1):
                p = doc.add_paragraph()
                set_rtl_and_justify(p)
                set_font_style(p.add_run(f"{idx}. "), size=12, bold=True)
                set_font_style(p.add_run(str(item).strip()), size=12)
            doc.add_paragraph(" ")

    doc.save(path)
    return path